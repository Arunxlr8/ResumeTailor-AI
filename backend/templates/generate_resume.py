"""
generate_resume.py
-------------------
Parametric resume builder using python-docx. Recreates exact document formatting,
two-column layout grid, margins, borders, fonts, and compact 2-page spacing.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

GREY = "ACA8AA"
BLACK = "000000"
FONT = "Times New Roman"

DEFAULT_CONTEXT = {
    "name": "ARUN KUMAR S",
    "title": "GENAI ENGINEER",
    "linkedin_url": "http://www.linkedin.com/in/arun-s-ak",
    "linkedin_text": "www.linkedin.com/in/arun-s-ak",
    "email_url": "mailto:arunduxq@gmail.com",
    "email_text": "arunduxq@gmail.com",
    "phone": "+91-9080243915",
    "summary": "GenAI Engineer with 3+ years of hands-on experience specialising in Large Language Models, Agentic AI Architectures, RAG pipelines, and LangGraph-based multi-stage workflows. Expert in designing intelligent, scalable AI automation systems using Google ADK, LangChain, and multi-agent orchestration frameworks. Proven expertise in DS, ML, GenAI, and Agentic AI delivering production-grade solutions across automotive, aerospace, and SDLC domains. Strong foundation in Python, FastAPI, Flask, and React enabling full-stack deployment of production-ready AI engineering solutions.",
    "skills": [
        ("Agentic AI & Frameworks: ", "Google ADK (Agent Development Kit), LangChain, LangGraph, Multi-Agent Orchestration, HITL Workflows, Custom Tools, MCP (Model Context Protocol), Orchestrator-Agent Pattern"),
        ("GenAI & LLM: ", "LLM, RAG, Prompt Engineering (Few-Shot, System Prompts), Structured Outputs, Embeddings, Semantic Search, Hallucination Detection, Evaluation Frameworks, Ollama, vLLM, Hugging Face"),
        ("DS & ML: ", "Scikit-learn, Pandas, NumPy, Similarity Search, Defect Classification, Cosine Similarity, Statistical Validation (SciPy, ttest_rel), Data Preprocessing, Anomaly Detection"),
        ("Vector DBs & RAG: ", "ChromaDB, FAISS, Vector Databases, Custom Chunking, RAG Document Indexing, Semantic Retrieval, Hybrid Retrieval, Embedding Models"),
        ("Programming & Frameworks: ", "Python (OOP), FastAPI, Flask, React, PyTest, SQLite, REST APIs, Selenium, JIRA Integration"),
        ("DevOps & Tools: ", "Docker, Git, GitHub, Jenkins, Azure DevOps, LM Studio, ECU Test (Tracetronics), Object API (RPC), VectorCast, CANoe")
    ],
    "experience": [
        {
            "role": "AI ENGINEER",
            "company_location": "L&T TECHNOLOGY SERVICES | MYSORE",
            "date": "2023 -- Present",
            "bullets": [
                "Designed and deployed multi-stage LLM orchestration pipelines and agentic AI systems for automotive ECU validation, aerospace verification, and IVI automation across Tier-1 global OEM clients.",
                "Built RAG-backed AI chatbots and domain assistants using LangChain, ChromaDB, FAISS, and vLLM for automated test case generation, requirement analysis, and log analysis workflows.",
                "Architected LangGraph-based multi-agent workflows with HITL (Human-In-The-Loop) interaction, session-wise memory management, and iterative feedback loops for aerospace test generation.",
                "Integrated AI workflows with domain-specific tools via RPC and Object API layers, converting LLM-generated structured outputs into executable test scripts and validation packages.",
                "Delivered full-stack AI applications using Flask, FastAPI, and React providing sandbox interfaces for automotive and aerospace validation engineers to interact with AI solutions.",
                "Implemented desktop automation scripts using Python OOP and open-source libraries to eliminate repetitive manual test execution in VectorCast and CANoe environments."
            ]
        }
    ],
    "projects": [
        {
            "title": "ECU Package Generation -- IVI & Powertrain Domain",
            "lines": [
                "Built an AI-driven workflow to auto-generate ECU-compatible test packages from High Level Test Cases using LLM-based metadata extraction, RAG-backed YAML schema generation, and Rule-Based Engine orchestration via Tracetronics Object API. Deployed in production for a Tier-1 German automotive OEM."
            ],
            "stack": "Python, LangChain, ChromaDB, Mistral AI, RAG, Object API (RPC), Flask, React"
        },
        {
            "title": "AI-Powered SDLC Test Lifecycle Automation",
            "intro": "Complete Agentic AI platform automating the end-to-end software testing lifecycle from requirement ingestion to defect analysis. Built using Google ADK with four specialized AI agents connected by a central Orchestrator Agent:",
            "bullets": [
                "AI1 -- Test Case Generation Agent: Two sub-agents Normalizer sub-agent pre-processes and standardises raw requirements; Test Step Generation sub-agent generates structured test cases per requirement. Evaluation Agent validates output across 5 dimensions: Accuracy, Coverage, Faithfulness, Hallucination Detection, and Relevance. Test cases reviewable and pushable directly to JIRA.",
                "AI2 -- Test Script Generation Agent: Two sub-agents Mapping sub-agent maps test steps to corresponding common test libraries; YAML Generation sub-agent creates structured requirement schema with parameters for ECU Test script generation. Feeds into suite pipeline for automated script creation.",
                "AI3 -- Log Analyzer Agent: Analyzer sub-agent ingests executed test log reports, performs Root Cause Analysis, and generates an Observation Dashboard pinpointing which step failed and how to fix it.",
                "AI4 -- Defect Analyzer Agent: Similarity sub-agent checks proposed defects against existing JIRA defect database, returns confidence score and explanation to determine if defect is new or duplicate enabling informed defect creation."
            ],
            "stack": "Python, Google ADK, ChromaDB, Custom Tools, MCP, Embedding Models, FastAPI, vLLM, React, JIRA Integration"
        },
        {
            "title": "Aero MCDC Test Set Generator -- Aerospace Domain",
            "lines": [
                "Developed a LangGraph-based multi-agent system with HITL interaction and SQLite memory management for aerospace MCDC test generation. Pipeline stages include requirement validation, base expression derivation, variable extraction, data range extraction, MCDC truth table generation, and test set expansion with user feedback interrupts at each stage."
            ],
            "stack": "Python, LangChain, LangGraph, ChromaDB, LLM, SQLite Memory, Session Management, Flask API"
        },
        {
            "title": "AI-Powered Chatbots -- Multiple OEM Clients",
            "lines": [
                "Provisioned multiple domain-specific AI chatbots for automotive OEM clients covering log analysis, image processing, and requirement analysis use cases. Each assistant powered by RAG pipelines with ChromaDB or FAISS and deployed via vLLM for scalable inference."
            ],
            "stack": "Python, LangChain, ChromaDB, FAISS, LLM, vLLM, RAG"
        },
        {
            "title": "AIOS -- AI Intelligent Kernel System",
            "lines": [
                "Developed an AI-native kernel runtime for LLMs and agents with predictive resource allocation, eBPF telemetry, and adaptive scheduling for dynamic CPU/GPU optimization. Recognized with Special Jury Award at AIOS 2025 hackathon."
            ],
            "stack": "Python, eBPF, Adaptive Scheduling, Resource Management"
        }
    ],
    "achievements": [
        "Hackathon Winner -- Mysore Open Hack 2024: 1st place among 120+ teams with an AI-powered code review and improvement assistant.",
        "Special Jury Award 2025 -- Mysore Open Hack 2024 -- AIOS: Recognised for developing an AI-native kernel runtime concept during a corporate hackathon.",
        "Star of the Month -- Production Impact: ECU Package Generation system deployed in production for a Tier-1 German automotive OEM, eliminating manual test authoring."
    ],
    "education": [
        {
            "degree": "Bachelor of Technology (B.Tech) -- Information Technology",
            "institution": "Rajalakshmi Engineering College, Chennai, Tamil Nadu",
            "details": "June 2023 | CGPA: 7.85"
        }
    ],
    "strengths": [
        "Deep specialization in LLM orchestration, agentic AI frameworks, LangGraph pipelines, and RAG-backed automotive validation systems.",
        "Proven ability to integrate AI reasoning with domain-specific tools via API and RPC layers for production-ready automotive solutions.",
        "Strong cross-domain adaptability across GenAI, embedded systems, test automation, and full-stack AI application delivery.",
        "Focused on scalable, production-ready AI design with measurable reduction in manual engineering effort."
    ]
}


# ---------------------------------------------------------------- helpers --

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge_name, edge in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        if edge is None:
            continue
        tag = f'w:{edge_name}'
        el = tcBorders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            tcBorders.append(el)
        el.set(qn('w:val'), edge.get('val', 'single'))
        el.set(qn('w:sz'), str(edge.get('sz', 4)))
        el.set(qn('w:space'), str(edge.get('space', 0)))
        el.set(qn('w:color'), edge.get('color', 'auto'))


def set_cell_margins(cell, top=0, left=0, bottom=0, right=0):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        mar.append(node)
    tcPr.append(mar)


def set_vertical_alignment(cell, align="bottom"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    va = OxmlElement('w:vAlign')
    va.set(qn('w:val'), align)
    tcPr.append(va)


def contextual_spacing(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(OxmlElement('w:contextualSpacing'))


def set_spacing(paragraph, before=None, after=None, line=None):
    pf = paragraph.paragraph_format
    if before is not None:
        pf.space_before = Twips(before)
    if after is not None:
        pf.space_after = Twips(after)
    if line is not None:
        pf.line_spacing = line / 240.0


def add_run(paragraph, text, bold=False, italic=False, size=10, color=BLACK, font=FONT):
    r = paragraph.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = RGBColor.from_string(color)
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is not None:
        rFonts.set(qn('w:eastAsia'), font)
    return r


def add_hyperlink(paragraph, url, text, size=8, color=BLACK, font=FONT, underline=False):
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    rFonts.set(qn('w:cs'), font)
    rPr.append(rFonts)

    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), color)
    color_el.set(qn('w:themeColor'), 'text1')
    rPr.append(color_el)

    if not underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'none')
        rPr.append(u)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(size * 2))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(size * 2))
    rPr.append(szCs)

    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_bullet(cell, text_runs, before=24, after=24, line=240, indent=440, hanging=260):
    p = cell.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(indent))
    ind.set(qn('w:hanging'), str(hanging))
    pPr.append(ind)
    contextual_spacing(p)
    set_spacing(p, before=before, after=after, line=line)
    add_run(p, "\u2022\u2002", size=10)
    for text, bold in text_runs:
        add_run(p, text, bold=bold, size=10)
    return p


def right_tab_paragraph(cell, pos=10700):
    p = cell.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(pos))
    tabs.append(tab)
    pPr.append(tabs)
    return p


def add_section_row(table, title):
    row = table.add_row()
    row.cells[0].merge(row.cells[1])
    cell = row.cells[0]
    set_cell_margins(cell, top=30, left=0, bottom=30, right=0)
    set_cell_border(cell,
                    top={'val': 'single', 'sz': 20, 'color': GREY},
                    bottom={'val': 'single', 'sz': 20, 'color': GREY},
                    left={'val': 'none'}, right={'val': 'none'})
    p = cell.paragraphs[0]
    add_run(p, title, bold=True, size=16)
    return cell


def add_body_row(table, top_border=False):
    row = table.add_row()
    row.cells[0].merge(row.cells[1])
    cell = row.cells[0]
    set_cell_margins(cell, top=40, left=0, bottom=40, right=0)
    set_cell_border(cell,
                    top={'val': 'single', 'sz': 20, 'color': GREY} if top_border else {'val': 'none'},
                    bottom={'val': 'single', 'sz': 20, 'color': GREY},
                    left={'val': 'none'}, right={'val': 'none'})
    cell.paragraphs[0].text = ""
    return cell


def generate_resume(context: dict = None, output_path: str = "Arun_Kumar_S_Resume_recreated.docx") -> str:
    """Generate DOCX resume using parametric context data."""
    ctx = copy.deepcopy(DEFAULT_CONTEXT)
    if context:
        for k, v in context.items():
            if v:
                ctx[k] = v

    doc = Document()

    # ---- page setup ----
    section = doc.sections[0]
    section.page_width = Twips(12240)
    section.page_height = Twips(15840)
    section.top_margin = Twips(450)
    section.bottom_margin = Twips(450)
    section.left_margin = Twips(720)
    section.right_margin = Twips(720)
    section.header_distance = Twips(708)
    section.footer_distance = Twips(708)

    # ---- document default font ----
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    rpr = normal.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT)

    # ---- whole-page 2-column layout table ----
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for col in table.columns:
        for c in col.cells:
            c.width = Twips(5400)

    header_row = table.rows[0]
    left_cell, right_cell = header_row.cells

    # left cell: name + title
    set_cell_margins(left_cell, top=60, left=0, bottom=60, right=40)
    set_cell_border(left_cell, top={'val': 'none'}, left={'val': 'none'},
                    right={'val': 'none'}, bottom={'val': 'single', 'sz': 20, 'color': GREY})
    set_vertical_alignment(left_cell, 'bottom')

    p_name = left_cell.paragraphs[0]
    contextual_spacing(p_name)
    set_spacing(p_name, after=20)
    add_run(p_name, ctx.get("name", "ARUN KUMAR S"), bold=True, size=16)

    p_blank = left_cell.add_paragraph()
    contextual_spacing(p_blank)
    set_spacing(p_blank, after=20)

    p_title = left_cell.add_paragraph()
    contextual_spacing(p_title)
    add_run(p_title, ctx.get("title", "GENAI ENGINEER"), bold=True, size=12)

    # right cell: contact block
    set_cell_margins(right_cell, top=60, left=1400, bottom=60, right=0)
    set_cell_border(right_cell, top={'val': 'none'}, left={'val': 'none'},
                    right={'val': 'none'}, bottom={'val': 'single', 'sz': 20, 'color': GREY})
    set_vertical_alignment(right_cell, 'bottom')

    p_spacer = right_cell.paragraphs[0]
    add_run(p_spacer, "", size=8)

    if ctx.get("linkedin_text") and ctx.get("linkedin_url"):
        p_linkedin = right_cell.add_paragraph()
        contextual_spacing(p_linkedin)
        add_hyperlink(p_linkedin, ctx["linkedin_url"], ctx["linkedin_text"], size=8)

    if ctx.get("email_text") and ctx.get("email_url"):
        p_email = right_cell.add_paragraph()
        contextual_spacing(p_email)
        add_hyperlink(p_email, ctx["email_url"], ctx["email_text"], size=8)

    if ctx.get("phone"):
        p_phone = right_cell.add_paragraph()
        contextual_spacing(p_phone)
        set_spacing(p_phone, after=14)
        add_run(p_phone, ctx["phone"], size=9.5)

    # ---- SUMMARY ----
    cell = add_body_row(table, top_border=True)
    p = cell.paragraphs[0]
    set_spacing(p, before=40, after=40, line=240)
    add_run(p, ctx.get("summary", ""), size=10)

    # ---- SKILLS ----
    add_section_row(table, "SKILLS")
    cell = add_body_row(table)
    cell.paragraphs[0].text = ""

    skills = ctx.get("skills", [])
    if isinstance(skills, list):
        for item in skills:
            if isinstance(item, (tuple, list)) and len(item) == 2:
                label, rest = item
                add_bullet(cell, [(str(label), True), (str(rest), False)], before=20, after=20, line=240)
            elif isinstance(item, dict):
                label = item.get("category", "")
                rest = item.get("items", "")
                add_bullet(cell, [(f"{label}: " if label else "", True), (str(rest), False)], before=20, after=20, line=240)
            elif isinstance(item, str):
                add_bullet(cell, [(item, False)], before=20, after=20, line=240)

    # ---- EXPERIENCE ----
    add_section_row(table, "EXPERIENCE")
    cell = add_body_row(table)
    cell.paragraphs[0].text = ""

    experiences = ctx.get("experience", [])
    for exp in experiences:
        p_job = right_tab_paragraph(cell)
        set_spacing(p_job, before=40, after=20)
        add_run(p_job, f"{exp.get('role', '')} ", bold=True, size=10)
        add_run(p_job, exp.get('company_location', ''), size=10)
        p_job.add_run("\t")
        add_run(p_job, exp.get('date', ''), italic=True, size=9)

        for b in exp.get("bullets", []):
            add_bullet(cell, [(b, False)], before=20, after=20, line=240)

    # ---- PROJECTS ----
    add_section_row(table, "PROJECTS")
    cell = add_body_row(table)
    cell.paragraphs[0].text = ""

    projects = ctx.get("projects", [])
    for prj in projects:
        p_title = cell.add_paragraph()
        set_spacing(p_title, before=50, after=12)
        add_run(p_title, prj.get("title", ""), bold=True, size=10)

        if prj.get("intro"):
            p_intro = cell.add_paragraph()
            set_spacing(p_intro, before=8, after=14, line=240)
            add_run(p_intro, prj["intro"], size=10)

        for line in prj.get("lines", []):
            p_line = cell.add_paragraph()
            pPr = p_line._p.get_or_add_pPr()
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), '160')
            pPr.append(ind)
            set_spacing(p_line, before=0, after=8)
            add_run(p_line, line, size=10)

        for bullet in prj.get("bullets", []):
            add_bullet(cell, [(bullet, False)], before=18, after=18, line=240)

        if prj.get("stack"):
            p_stack = cell.add_paragraph()
            pPr = p_stack._p.get_or_add_pPr()
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), '160')
            pPr.append(ind)
            set_spacing(p_stack, after=36)
            add_run(p_stack, "Stack: ", bold=True, size=10)
            add_run(p_stack, prj["stack"], size=10)

    # ---- ACHIEVEMENTS ----
    add_section_row(table, "ACHIEVEMENTS")
    cell = add_body_row(table)
    cell.paragraphs[0].text = ""
    for ach in ctx.get("achievements", []):
        add_bullet(cell, [(ach, False)], before=20, after=20, line=240)

    # ---- EDUCATION ----
    add_section_row(table, "EDUCATION")
    cell = add_body_row(table)
    cell.paragraphs[0].text = ""

    education_list = ctx.get("education", [])
    for edu in education_list:
        p_deg = cell.paragraphs[0] if cell.paragraphs[0].text == "" else cell.add_paragraph()
        set_spacing(p_deg, before=40, after=14)
        add_run(p_deg, edu.get("degree", ""), bold=True, size=10)

        p_col = cell.add_paragraph()
        set_spacing(p_col, after=12)
        add_run(p_col, edu.get("institution", ""), size=10)

        p_gpa = cell.add_paragraph()
        set_spacing(p_gpa, after=36)
        add_run(p_gpa, edu.get("details", ""), bold=True, size=10)

    # ---- STRENGTHS ----
    add_section_row(table, "STRENGTHS")
    row = table.add_row()
    row.cells[0].merge(row.cells[1])
    cell = row.cells[0]
    set_cell_margins(cell, top=30, left=0, bottom=30, right=0)
    set_cell_border(cell, top={'val': 'single', 'sz': 20, 'color': GREY},
                    bottom={'val': 'none'}, left={'val': 'none'}, right={'val': 'none'})
    cell.paragraphs[0].text = ""
    for strg in ctx.get("strengths", []):
        add_bullet(cell, [(strg, False)], before=20, after=20, line=240)

    # Save output
    out_file = Path(output_path)
    out_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_file))
    return str(out_file)


if __name__ == "__main__":
    out = generate_resume()
    print(f"Successfully generated resume: {out}")
