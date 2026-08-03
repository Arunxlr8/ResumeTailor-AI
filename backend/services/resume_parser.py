"""Resume parsing service.

Detects resume file type (DOCX, PDF, TXT) automatically by extension and file signature,
then extracts all text content for parsing and tailoring by downstream agent nodes.
"""

from pathlib import Path
from docx import Document
import fitz  # PyMuPDF
from utils.logging import get_thread_logger


def detect_file_format(path: str) -> str:
    """Detect file format based on suffix and/or magic bytes.

    Checks the file extension first. If it is empty or unrecognized, falls back
    to reading file magic bytes (PDF signature, DOCX PK header).

    Parameters:
        path (str): The file path to inspect.

    Returns:
        str: The lowercase extension (e.g. '.docx', '.pdf', '.txt').
    """
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix in [".docx", ".pdf", ".txt"]:
        return suffix

    # Fallback to inspecting magic bytes
    try:
        with open(path, "rb") as f:
            header = f.read(4)
        if header.startswith(b"%PDF"):
            return ".pdf"
        if header.startswith(b"PK\x03\x04"):
            return ".docx"
    except Exception:
        pass

    return ".txt"


def parse_docx(path: str) -> str:
    """Extract and combine all text from a Word DOCX document.

    Extracts text from paragraphs and table cells, discarding empty sections.

    Parameters:
        path (str): Filepath to the DOCX document.

    Returns:
        str: Extracted and cleaned text content.
    """
    document = Document(path)
    text = []
    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        if value:
            text.append(value)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                value = cell.text.strip()
                if value:
                    text.append(value)
    return "\n".join(text)


def parse_pdf(path: str) -> str:
    """Extract and combine text from all pages in a PDF document.

    Parameters:
        path (str): Filepath to the PDF document.

    Returns:
        str: Extracted and cleaned text content.
    """
    document = fitz.open(path)
    text = []
    try:
        for page in document:
            value = page.get_text().strip()
            if value:
                text.append(value)
    finally:
        document.close()
    return "\n".join(text)


def parse_resume(path: str, thread_id: str | None = None) -> str:
    """Automatically detect and extract text from DOCX, PDF, or TXT resume files.

    Logs execution progress and raises detailed errors if parsing fails.

    Parameters:
        path (str): The filepath of the resume to parse.
        thread_id (str | None): Optional thread ID to log progress and errors to thread log.

    Returns:
        str: Cleaned text representation of the resume.

    Raises:
        ValueError: If file format is unsupported or parsing fails.
    """
    logger = get_thread_logger(thread_id)
    logger.info(f"Detecting format and parsing resume: {path}")

    ext = detect_file_format(path)
    logger.info(f"Detected file format: {ext}")

    if ext == ".docx":
        try:
            content = parse_docx(path)
            logger.info("Successfully parsed DOCX resume.")
            return content
        except Exception as e:
            error_msg = f"Failed to parse DOCX resume: {str(e)}"
            logger.error(error_msg)
            raise ValueError(error_msg) from e

    if ext == ".pdf":
        try:
            content = parse_pdf(path)
            logger.info("Successfully parsed PDF resume.")
            return content
        except Exception as e:
            error_msg = f"Failed to parse PDF resume: {str(e)}"
            logger.error(error_msg)
            raise ValueError(error_msg) from e

    if ext == ".txt":
        try:
            with open(path, "r", encoding="utf-8") as file:
                content = file.read()
                logger.info("Successfully parsed UTF-8 text resume.")
                return content
        except UnicodeDecodeError:
            try:
                # Fallback to Latin-1 encoding
                with open(path, "r", encoding="latin-1") as file:
                    content = file.read()
                    logger.info("Successfully parsed Latin-1 text resume.")
                    return content
            except Exception as e:
                error_msg = f"Failed to read text resume with Latin-1 fallback: {str(e)}"
                logger.error(error_msg)
                raise ValueError(error_msg) from e
        except Exception as e:
            error_msg = f"Failed to read text resume: {str(e)}"
            logger.error(error_msg)
            raise ValueError(error_msg) from e

    error_msg = f"Unsupported resume format: {ext} for file: {path}"
    logger.error(error_msg)
    raise ValueError(error_msg)